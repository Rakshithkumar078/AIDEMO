import PyPDF2
from docx import Document as DocxDocument
from unstructured.partition.auto import partition
import io
from typing import List, Dict, Any

async def process_document(file_content: bytes, filename: str) -> List[Dict[str, Any]]:
    """Process document and return chunks for embedding"""
    
    file_extension = filename.lower().split('.')[-1]
    
    if file_extension == 'pdf':
        return await process_pdf(file_content)
    elif file_extension == 'docx':
        return await process_docx(file_content)
    elif file_extension == 'txt':
        return await process_text(file_content)
    else:
        # Use unstructured for other formats
        return await process_with_unstructured(file_content)

async def extract_text_content(file_content: bytes, filename: str) -> str:
    """Extract raw text content from document for viewing"""
    
    file_extension = filename.lower().split('.')[-1]
    
    try:
        if file_extension == 'pdf':
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        elif file_extension == 'docx':
            doc = DocxDocument(io.BytesIO(file_content))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        elif file_extension == 'txt':
            return file_content.decode('utf-8')
        elif file_extension in ['jpg', 'jpeg', 'png', 'gif']:
            return f"[Image file: {filename}]\nImage content cannot be displayed as text."
        else:
            # Try to decode as text
            try:
                return file_content.decode('utf-8')
            except:
                return f"[Binary file: {filename}]\nContent cannot be displayed as text."
    except Exception as e:
        return f"Error reading file content: {str(e)}"

async def process_pdf(file_content: bytes) -> List[Dict[str, Any]]:
    """Extract text from PDF and chunk it"""
    pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
    text = ""
    
    for page in pdf_reader.pages:
        text += page.extract_text() + "\n"
    
    return chunk_text(text)

async def process_docx(file_content: bytes) -> List[Dict[str, Any]]:
    """Extract text from DOCX and chunk it"""
    doc = DocxDocument(io.BytesIO(file_content))
    text = ""
    
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    
    return chunk_text(text)

async def process_text(file_content: bytes) -> List[Dict[str, Any]]:
    """Process plain text file"""
    text = file_content.decode('utf-8')
    return chunk_text(text)

async def process_with_unstructured(file_content: bytes) -> List[Dict[str, Any]]:
    """Use unstructured library for other formats"""
    try:
        elements = partition(file=io.BytesIO(file_content))
        text = "\n".join([str(element) for element in elements])
        return chunk_text(text)
    except Exception:
        # Fallback to treating as text
        try:
            text = file_content.decode('utf-8')
            return chunk_text(text)
        except:
            return [{"content": "Unable to process file", "metadata": {"error": True}}]

def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[Dict[str, Any]]:
    """Split text into overlapping chunks"""
    chunks = []
    start = 0
    
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        
        chunks.append({
            "content": chunk.strip(),
            "metadata": {
                "start": start,
                "end": end,
                "length": len(chunk.strip())
            }
        })
        
        start = end - overlap
        if start >= len(text):
            break
    
    return chunks